"""
RAG System for HealthYoda Knowledge Base
Extracts clinical interview patterns from text files, provides semantic retrieval using vector database
Supports flexible schema for clinical patterns with priority-based retrieval
"""
import docx
import os
from typing import List, Dict, Optional, Tuple
import json
from text_processor import TextFileProcessor
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import chromadb
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    logger.warning("ChromaDB not installed - vector database features disabled")

class QuestionBookRAG:
    """RAG system for retrieving relevant clinical interview patterns"""

    def __init__(self, docx_path: str = 'docx/Question BOOK.docx', txt_folder: str = 'txt', openai_client=None, model_manager=None):
        self.docx_path = docx_path
        self.txt_folder = txt_folder
        self.questions = []  # Stores all patterns (from both docx and txt files)
        self.sections = []
        self.text_processor = TextFileProcessor(txt_folder)

        # Handle different client types
        if model_manager and hasattr(model_manager, 'create_embeddings'):
            # New model manager interface
            self.model_manager = model_manager
            self.openai_client = None
        elif openai_client and hasattr(openai_client, 'create_embeddings'):
            # Legacy OpenAI client
            self.openai_client = openai_client
            self.model_manager = None
        elif openai_client and hasattr(openai_client, 'create_chat_completion'):
            # Model manager passed as openai_client parameter
            self.model_manager = openai_client
            self.openai_client = None
        else:
            self.openai_client = None
            self.model_manager = None
        
        # Initialize ChromaDB for vector storage (if available)
        self.chroma_client = None
        self.collection = None
        
        if CHROMADB_AVAILABLE:
            try:
                # Use new ChromaDB client API
                self.chroma_client = chromadb.PersistentClient(path="./chroma_db")
                
                # Create or get collection
                self.collection = self.chroma_client.get_or_create_collection(
                    name="question_book",
                    metadata={"description": "HealthYoda Clinical Interview Patterns"}
                )
            except Exception as e:
                logger.error(f"ChromaDB initialization failed: {e}")
                import traceback
                traceback.print_exc()
                self.chroma_client = None
                self.collection = None
        
        # Check if rebuild is requested via environment variable
        rebuild_vectorstore = os.getenv('REBUILD_VECTORSTORE', 'false').lower() == 'true'
        
        # Check if vectorstore already has data
        existing_count = self.collection.count() if self.collection else 0
        
        if rebuild_vectorstore:
            # Rebuild explicitly requested - always extract and rebuild
            if existing_count > 0:
                logger.info(f"[REBUILD] Rebuild requested - extracting documents and rebuilding vectorstore (current: {existing_count} embeddings)...")
            else:
                logger.info("[REBUILD] Rebuild requested - extracting documents and building vectorstore...")
            
            # Load documents (docx and text files)
            self.load_documents()
            
            # Create embeddings with incremental rebuild
            if self.collection:
                self.create_embeddings(force_rebuild=True)
        
        elif existing_count == 0:
            # First run - vectorstore is empty, need to build it
            logger.info("[INIT] First run detected - loading documents and building vectorstore...")
            
            # Load documents (docx and text files)
            self.load_documents()
            
            # Create embeddings
            if self.collection:
                self.create_embeddings(force_rebuild=False)
        
        else:
            # Vectorstore exists and no rebuild requested - load metadata only (no extraction)
            logger.info(f"[OK] Using existing vectorstore with {existing_count} embeddings")
            logger.info(f"[OK] Skipping document extraction to save tokens (set REBUILD_VECTORSTORE=true to rebuild)")
            
            # Load lightweight metadata from vectorstore for stats and fallback search
            self._load_metadata_from_vectorstore()
            
            # Log available domains and content types
            self._log_available_patterns()
    
    def _load_metadata_from_vectorstore(self):
        """Load lightweight pattern metadata from vectorstore (no extraction needed)"""
        if not self.collection:
            self.questions = []
            return
        
        try:
            # Get all metadata from vectorstore (lightweight query)
            results = self.collection.get(include=['metadatas'])
            
            if not results or not results.get('metadatas'):
                logger.warning("[WARNING] No metadata found in vectorstore")
                self.questions = []
                return
            
            # Create lightweight pattern objects from metadata
            self.questions = []
            for metadata in results['metadatas']:
                # Parse tags from comma-separated string
                tags_str = metadata.get('tags', '')
                tags = tags_str.split(',') if tags_str else []
                
                # Create minimal pattern object for search compatibility
                pattern = {
                    'medical_domain': metadata.get('medical_domain', ''),
                    'section': metadata.get('section', ''),
                    'content_type': metadata.get('content_type', ''),
                    'priority': metadata.get('priority', 'NORMAL'),
                    'tree_path': metadata.get('tree_path', ''),
                    'source': metadata.get('source', ''),
                    'bot_question': '',  # Not stored in metadata to save space
                    'clinical_context': '',
                    'expected_patient_responses': [],
                    'red_flags': [],
                    'tags': tags
                }
                self.questions.append(pattern)
            
            logger.info(f"[OK] Loaded metadata for {len(self.questions)} patterns from vectorstore")
            
        except Exception as e:
            logger.error(f"Error loading metadata from vectorstore: {e}")
            self.questions = []
    
    def _log_available_patterns(self):
        """Log available patterns summary for debugging"""
        if not self.questions:
            logger.warning("[WARNING] No patterns available")
            return
        
        # Count by domain
        domains = {}
        for q in self.questions:
            domain = q.get('medical_domain', 'Unknown')
            domains[domain] = domains.get(domain, 0) + 1
        
        # Count by content type
        content_types = {}
        for q in self.questions:
            ctype = q.get('content_type', 'unknown')
            content_types[ctype] = content_types.get(ctype, 0) + 1
        
        # Count by priority
        priorities = {}
        for q in self.questions:
            priority = q.get('priority', 'NORMAL')
            priorities[priority] = priorities.get(priority, 0) + 1
        
        logger.info(f"[RAG] Available Content Types: {dict(sorted(content_types.items(), key=lambda x: x[1], reverse=True))}")
        logger.info(f"[RAG] Priority Distribution: {dict(sorted(priorities.items(), key=lambda x: ['CRITICAL', 'HIGH', 'NORMAL', 'LOW'].index(x[0])))}")
        logger.info(f"[RAG] Medical Domains: {len(domains)} domains ({', '.join(list(domains.keys())[:5])}...)")
    
    def load_documents(self):
        """Load both docx document and text files"""
        # Load docx document (if exists)
        if os.path.exists(self.docx_path):
            self.load_docx_document()
        else:
            logger.warning(f"DOCX file not found at {self.docx_path}, skipping...")

        # Load text files (primary source)
        self.load_text_files()

    def load_docx_document(self):
        """Load and parse the .docx document (legacy format, maintained for backwards compatibility)"""
        if not os.path.exists(self.docx_path):
            logger.warning(f"Question book not found at {self.docx_path}")
            return
        
        logger.info(f"Loading legacy DOCX format from {self.docx_path}")
        
        doc = docx.Document(self.docx_path)
        current_system = None
        current_symptom = None
        current_category = None
        current_question = None
        current_answers = []
        
        # Categories that indicate question sections
        categories = ['Chief Complaint', 'Onset/Duration', 'Quality/Severity', 
                     'Aggravating/Relieving', 'Associated Symptoms', 'Red Flags', 
                     'ROS', 'Context']
        
        # Exclude patterns
        exclude_patterns = ['Table of Contents', 'HealthYoda History-Taking Handbook', 
                           'comprehensive system-wise', 'A comprehensive', '[page]']
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if not text:
                continue
            
            # Skip excluded patterns
            if any(pattern in text for pattern in exclude_patterns):
                continue
            
            # Detect system headers
            if 'HealthYoda History Framework' in text:
                # Extract system name
                if 'Cardiac' in text:
                    current_system = 'Cardiac System'
                elif 'Respiratory' in text:
                    current_system = 'Respiratory System'
                elif 'GI System' in text or ('GI' in text and 'System' in text):
                    current_system = 'GI System'
                elif 'Neurologic' in text:
                    current_system = 'Neurologic System'
                elif 'Musculoskeletal' in text:
                    current_system = 'Musculoskeletal System'
                elif 'GU System' in text or ('GU' in text and 'System' in text):
                    current_system = 'GU System'
                elif 'Dermatologic' in text:
                    current_system = 'Dermatologic System'
                elif 'Endocrine' in text or 'General' in text:
                    current_system = 'General/Endocrine/Infectious Disease System'
                elif 'ENT' in text or 'Eye' in text:
                    current_system = 'ENT/Eye System'
                
                if current_system:
                    current_symptom = None
                    current_category = None
                    current_question = None
                    current_answers = []
            
            # Detect symptom/complaint headers
            elif current_system and text not in categories and \
                 not text.startswith('Q') and not text.startswith('Possible') and \
                 not text.startswith('-') and len(text) < 100 and \
                 text != current_system and 'System' not in text and \
                 'HealthYoda' not in text and text != 'Wrap-up':
                # Save previous question
                if current_question and current_system:
                    self._save_question_legacy(current_system, current_symptom, current_category,
                                      current_question, current_answers.copy(), i)
                    current_question = None
                    current_answers = []
                
                current_symptom = text
                current_category = None
            
            # Detect category headers
            elif text in categories:
                # Save previous question
                if current_question and current_system:
                    self._save_question_legacy(current_system, current_symptom, current_category,
                                      current_question, current_answers.copy(), i)
                    current_question = None
                    current_answers = []
                
                current_category = text
            
            # Detect questions
            elif text.startswith('Q:') or text.startswith('Q.'):
                # Save previous question
                if current_question and current_system:
                    self._save_question_legacy(current_system, current_symptom, current_category,
                                      current_question, current_answers.copy(), i)
                
                current_question = text.replace('Q:', '').replace('Q.', '').strip()
                current_answers = []
            
            # Detect possible answers
            elif text.startswith('Possible Answers:') or text.startswith('Possible answers:'):
                current_answers = []
            
            elif text.startswith('-') and current_question:
                answer = text[1:].strip()
                if answer:
                    current_answers.append(answer)
        
        # Save last question
        if current_question and current_system:
            self._save_question_legacy(current_system, current_symptom, current_category,
                              current_question, current_answers.copy(), len(doc.paragraphs))
        
        systems_found = len(set(q.get('system') for q in self.questions if q.get('system')))
        logger.info(f"Loaded {len(self.questions)} legacy patterns from {systems_found} systems")

    def load_text_files(self):
        """Load and process text files from the txt folder using new flexible schema"""
        try:
            text_patterns = self.text_processor.extract_all_text_data()
            
            # Convert text patterns to unified format
            for pattern in text_patterns:
                # Map to unified schema (already in correct format from text_processor)
                self.questions.append(pattern)
            
            logger.info(f"Loaded {len(text_patterns)} clinical patterns from text files")
            
            # Log stats
            stats = self.text_processor.get_stats()
            logger.info(f"Text file stats: {stats.get('extraction_stats', {})}")
            
        except Exception as e:
            logger.error(f"Failed to load text files: {e}")
            import traceback
            traceback.print_exc()

    def _save_question_legacy(self, system, symptom, category, question, answers, line_number):
        """Helper method to save a legacy docx question with unified schema"""
        # Map legacy format to new flexible schema
        pattern = {
            'medical_domain': system,
            'section': symptom or 'General',
            'content_type': 'interview_question',
            'bot_question': question,
            'clinical_context': f"{system} - {symptom or 'General'} - {category or 'General'}",
            'expected_patient_responses': answers,
            'red_flags': [],
            'priority': 'HIGH' if category == 'Red Flags' else 'NORMAL',
            'tags': [
                f'clinical_domain:{system}',
                f'section:{symptom or "General"}',
                f'content_type:interview_question',
                f'priority:{"HIGH" if category == "Red Flags" else "NORMAL"}',
                'source:docx'
            ],
            'tree_path': " > ".join(filter(None, [system, symptom, category])),
            'source': 'docx',
            'metadata': {
                'line_number': line_number,
                'system': system,
                'symptom': symptom,
                'category': category
            }
        }
        
        self.questions.append(pattern)
    
    def create_embeddings(self, force_rebuild: bool = False, rebuild_model: str = None):
        """
        Create embeddings for all patterns and store in vector database
        Performs TRUE incremental rebuild - only adds new embeddings without clearing existing ones

        Args:
            force_rebuild: If True, rebuild embeddings (incremental - adds new ones only)
            rebuild_model: Specific model to use for embeddings (overrides default)
        """
        if not self.collection:
            logger.warning("Vector database not available - skipping embeddings. Using keyword search only.")
            return

        if not self.openai_client and not self.model_manager:
            logger.warning("No embedding client available - skipping embeddings. Using keyword search only.")
            return

        current_count = self.collection.count()

        # Check if collection already has data (unless force rebuild)
        if not force_rebuild and current_count > 0:
            logger.info(f"[OK] Vector database already has {current_count} embeddings")
            logger.info(f"   Set REBUILD_VECTORSTORE=true in .env to rebuild")
            return

        if not self.questions:
            logger.warning("No patterns to process")
            return

        # Incremental rebuild: only add NEW patterns
        if force_rebuild and current_count > 0:
            # Check if we have new patterns to add
            new_pattern_count = len(self.questions)
            
            if new_pattern_count > current_count:
                # We have new patterns - add only the new ones
                patterns_to_add = new_pattern_count - current_count
                logger.info(f"[INCREMENTAL] Found {patterns_to_add} new patterns to add to existing {current_count} embeddings")
                logger.info(f"[INCREMENTAL] Total after rebuild: {new_pattern_count} embeddings")
                
                # Only process the new patterns (from current_count onwards)
                questions_to_process = self.questions[current_count:]
                start_index = current_count
            elif new_pattern_count == current_count:
                logger.info(f"[OK] Vector store already up-to-date ({current_count} embeddings)")
                return
            else:
                # Fewer patterns than embeddings - something changed, do full rebuild
                logger.warning(f"[WARNING] Pattern count ({new_pattern_count}) < embedding count ({current_count})")
                logger.info(f"[REBUILD] Performing full rebuild due to pattern count mismatch...")
                
                # Clear and rebuild
                try:
                    self.chroma_client.delete_collection("question_book")
                    self.collection = self.chroma_client.create_collection(
                        name="question_book",
                        metadata={"description": "HealthYoda Clinical Interview Patterns"}
                    )
                    logger.info(f"[REBUILD] Collection cleared, rebuilding {len(self.questions)} embeddings...")
                except Exception as e:
                    logger.error(f"Error clearing collection: {e}")
                
                questions_to_process = self.questions
                start_index = 0
        else:
            # First build
            logger.info(f"[INIT] Creating embeddings for {len(self.questions)} patterns...")
            questions_to_process = self.questions
            start_index = 0

        logger.info(f"Processing {len(questions_to_process)} patterns (starting from index {start_index})...")

        # Create text for embedding (combine all relevant context)
        texts_to_embed = []
        ids = []
        metadatas = []

        for idx, pattern in enumerate(questions_to_process, start=start_index):
            # Create rich text for embedding: question + context + responses
            text_parts = []
            
            # Primary content
            if pattern.get('bot_question'):
                text_parts.append(f"Question: {pattern['bot_question']}")
            
            if pattern.get('clinical_context'):
                text_parts.append(f"Context: {pattern['clinical_context']}")
            
            # Domain and section
            if pattern.get('medical_domain'):
                text_parts.append(f"Domain: {pattern['medical_domain']}")
            
            if pattern.get('section'):
                text_parts.append(f"Section: {pattern['section']}")
            
            # Content type
            if pattern.get('content_type'):
                text_parts.append(f"Type: {pattern['content_type']}")
            
            # Expected responses
            if pattern.get('expected_patient_responses'):
                responses = ', '.join(pattern['expected_patient_responses'][:3])
                text_parts.append(f"Expected: {responses}")
            
            # Red flags (important for matching)
            if pattern.get('red_flags'):
                flags = ', '.join(pattern['red_flags'][:3])
                text_parts.append(f"Red Flags: {flags}")
            
            # Differentials
            if pattern.get('differentials'):
                diffs = ', '.join(pattern['differentials'][:5])
                text_parts.append(f"Differentials: {diffs}")

            full_text = " | ".join(text_parts)
            texts_to_embed.append(full_text)
            ids.append(f"q_{idx}")
            
            # Store metadata for filtering
            metadatas.append({
                'medical_domain': pattern.get('medical_domain', ''),
                'section': pattern.get('section', ''),
                'content_type': pattern.get('content_type', ''),
                'priority': pattern.get('priority', 'NORMAL'),
                'tree_path': pattern.get('tree_path', ''),
                'source': pattern.get('source', ''),
                'tags': ','.join(pattern.get('tags', [])),  # Store as comma-separated string
                'index': idx
            })
        
        # Create embeddings in batches
        batch_size = 100
        for i in range(0, len(texts_to_embed), batch_size):
            batch_texts = texts_to_embed[i:i+batch_size]
            batch_ids = ids[i:i+batch_size]
            batch_metadatas = metadatas[i:i+batch_size]
            
            try:
                # Get embeddings using specified rebuild model or fallback logic
                embeddings = None

                # If rebuild_model is specified, use it directly with OpenAI client
                if rebuild_model and self.openai_client:
                    try:
                        response = self.openai_client.embeddings.create(
                            model=rebuild_model,
                            input=batch_texts
                        )
                        embeddings = [item.embedding for item in response.data]
                        logger.info(f"  Using specified rebuild model: {rebuild_model}")
                    except Exception as e:
                        logger.warning(f"Specified rebuild model {rebuild_model} failed: {e}, trying fallback...")

                # Fallback: use model manager or direct OpenAI client
                if embeddings is None:
                    if self.model_manager and hasattr(self.model_manager, 'create_embeddings'):
                        try:
                            response = self.model_manager.create_embeddings(batch_texts)
                            embeddings = [item.embedding for item in response.data]
                        except Exception as e:
                            logger.warning(f"Model manager embeddings failed: {e}, trying direct client...")

                    if embeddings is None and self.openai_client:
                        response = self.openai_client.embeddings.create(
                            model="text-embedding-3-small",  # Cost-effective embedding model
                            input=batch_texts
                        )
                        embeddings = [item.embedding for item in response.data]

                if embeddings is None:
                    raise ValueError("No embedding client available")
                
                # Add to ChromaDB
                self.collection.add(
                    embeddings=embeddings,
                    documents=batch_texts,
                    ids=batch_ids,
                    metadatas=batch_metadatas
                )
                
                logger.info(f"  Processed {min(i+batch_size, len(texts_to_embed))}/{len(texts_to_embed)} patterns...")
            except Exception as e:
                logger.error(f"  Error creating embeddings for batch {i}: {e}")
                continue
        
        logger.info(f"[OK] Created {self.collection.count()} embeddings in vector database")
    
    def search_by_system(self, system_name: str) -> List[Dict]:
        """Retrieve all patterns for a specific system/domain"""
        return [q for q in self.questions if system_name.lower() in q.get('medical_domain', '').lower()]
    
    def search_by_symptom(self, symptom: str) -> List[Dict]:
        """Retrieve patterns for a specific symptom/section"""
        symptom_lower = symptom.lower()
        return [q for q in self.questions 
                if q.get('section') and symptom_lower in q['section'].lower()]
    
    def search_by_content_type(self, content_type: str) -> List[Dict]:
        """Retrieve patterns by content type (red_flag, differential, interview_question, clinical_clue)"""
        return [q for q in self.questions if q.get('content_type') == content_type]
    
    def search_by_priority(self, priority: str) -> List[Dict]:
        """Retrieve patterns by priority (CRITICAL, HIGH, NORMAL, LOW)"""
        return [q for q in self.questions if q.get('priority') == priority]
    
    def search_by_keywords(self, keywords: List[str], medical_domain: Optional[str] = None) -> List[Dict]:
        """Search patterns by keywords"""
        results = []
        keywords_lower = [k.lower() for k in keywords]
        
        for q in self.questions:
            if medical_domain and medical_domain.lower() not in q.get('medical_domain', '').lower():
                continue
            
            # Search in question, context, and responses
            searchable_text = ' '.join([
                q.get('bot_question', ''),
                q.get('clinical_context', ''),
                ' '.join(q.get('expected_patient_responses', []))
            ]).lower()
            
            if any(kw in searchable_text for kw in keywords_lower):
                results.append(q)
        
        return results
    
    def get_next_question(self, conversation_context: str, current_category: Optional[str] = None,
                         symptom: Optional[str] = None, system: Optional[str] = None, 
                         use_semantic_search: bool = True, prioritize_red_flags: bool = True) -> Optional[Dict]:
        """
        Get the next relevant question based on conversation context
        Uses semantic search (vector DB) with priority-based ordering
        
        Args:
            conversation_context: Current conversation context
            current_category: Current category filter
            symptom: Symptom filter
            system: Medical domain filter
            use_semantic_search: Whether to use semantic search (vs keyword)
            prioritize_red_flags: Whether to prioritize red flags in results
        
        Returns:
            Pattern dict with question and metadata, or None
        """
        # Try semantic search first if embeddings are available
        if use_semantic_search and self.collection and self.collection.count() > 0:
            try:
                # Create query embedding
                if self.model_manager and hasattr(self.model_manager, 'create_embeddings'):
                    try:
                        response = self.model_manager.create_embeddings([conversation_context])
                        query_embedding = response.data[0].embedding
                    except Exception:
                        if self.openai_client:
                            response = self.openai_client.embeddings.create(
                                model="text-embedding-3-small",
                                input=conversation_context
                            )
                            query_embedding = response.data[0].embedding
                        else:
                            raise ValueError("No embedding client available")
                elif self.openai_client:
                    response = self.openai_client.embeddings.create(
                        model="text-embedding-3-small",
                        input=conversation_context
                    )
                    query_embedding = response.data[0].embedding
                else:
                    raise ValueError("No embedding client available")
                
                # Build where clause for filtering
                where_clause = None
                if system or symptom:
                    where_clause = {}
                    if system:
                        where_clause['medical_domain'] = {'$contains': system}
                    if symptom:
                        where_clause['section'] = {'$contains': symptom}
                
                # Query vector database - get more results to allow priority sorting
                query_kwargs = {
                    'query_embeddings': [query_embedding],
                    'n_results': 20,  # Get more results for priority filtering
                    'include': ['metadatas', 'documents', 'distances']  # Include documents for bot_question extraction
                }
                if where_clause:
                    query_kwargs['where'] = where_clause
                
                results = self.collection.query(**query_kwargs)
                
                if results['ids'] and len(results['ids'][0]) > 0:
                    # Also get documents to extract bot_question (needed when metadata-only patterns loaded)
                    documents = results.get('documents', [[]])[0] if results.get('documents') else []
                    
                    # Get all matching patterns
                    matching_patterns = []
                    for result_idx, result_id in enumerate(results['ids'][0]):
                        idx = int(result_id.replace('q_', ''))
                        
                        # Get pattern from self.questions (may be metadata-only)
                        pattern = self.questions[idx].copy() if idx < len(self.questions) else {}
                        
                        # If bot_question is empty (metadata-only), extract from document
                        if not pattern.get('bot_question') and result_idx < len(documents):
                            # Document format: "Question: <text> | Context: <text> | ..."
                            doc = documents[result_idx]
                            if doc and "Question: " in doc:
                                # Extract question from document
                                question_part = doc.split(" | ")[0]  # Get first part
                                if question_part.startswith("Question: "):
                                    pattern['bot_question'] = question_part.replace("Question: ", "").strip()
                        
                        pattern['_similarity_score'] = results.get('distances', [[]])[0][result_idx] if results.get('distances') else 0
                        matching_patterns.append(pattern)
                    
                    # Priority-based sorting if enabled
                    if prioritize_red_flags and matching_patterns:
                        # Sort by priority: CRITICAL > HIGH > NORMAL > LOW
                        priority_order = {'CRITICAL': 0, 'HIGH': 1, 'NORMAL': 2, 'LOW': 3}
                        matching_patterns.sort(key=lambda x: (
                            priority_order.get(x.get('priority', 'NORMAL'), 4),
                            x.get('_similarity_score', 1.0)  # Then by similarity
                        ))
                    
                    # Return top result
                    top_pattern = matching_patterns[0]
                    
                    # Log results
                    logger.info(f"[RAG] Semantic search found {len(matching_patterns)} relevant patterns")
                    logger.info(f"[RAG] Top match: {top_pattern.get('content_type')} - Priority: {top_pattern.get('priority')}")
                    if results.get('distances'):
                        logger.info(f"[RAG] Similarity score: {top_pattern.get('_similarity_score', 0):.4f}")
                    
                    return top_pattern
                    
            except Exception as e:
                logger.warning(f"[RAG] Semantic search failed: {e}, falling back to keyword search")
                import traceback
                traceback.print_exc()
        
        # Fallback to keyword-based search with priority ordering
        context_lower = conversation_context.lower()
        
        # Filter patterns
        candidates = self.questions
        
        if system:
            candidates = [q for q in candidates if system.lower() in q.get('medical_domain', '').lower()]
        
        if symptom:
            candidates = [q for q in candidates if q.get('section') and symptom.lower() in q['section'].lower()]
        
        # Keyword matching
        if candidates:
            scored_candidates = []
            for q in candidates:
                searchable = ' '.join([
                    q.get('bot_question', ''),
                    q.get('clinical_context', ''),
                    q.get('section', ''),
                    q.get('medical_domain', '')
                ]).lower()
                
                # Simple keyword matching score
                score = sum(1 for word in context_lower.split() if word in searchable and len(word) > 3)
                if score > 0:
                    scored_candidates.append((q, score))
            
            if scored_candidates:
                # Sort by priority, then by keyword score
                if prioritize_red_flags:
                    priority_order = {'CRITICAL': 0, 'HIGH': 1, 'NORMAL': 2, 'LOW': 3}
                    scored_candidates.sort(key=lambda x: (
                        priority_order.get(x[0].get('priority', 'NORMAL'), 4),
                        -x[1]  # Negative for descending score
                    ))
                else:
                    scored_candidates.sort(key=lambda x: -x[1])
                
                return scored_candidates[0][0]
        
        return None
    
    def get_all_categories_for_symptom(self, symptom: str) -> List[str]:
        """Get all content types available for a symptom"""
        patterns = self.search_by_symptom(symptom)
        content_types = list(set(q.get('content_type') for q in patterns if q.get('content_type')))
        return content_types
    
    def get_questions_by_phase(self, phase: str) -> List[Dict]:
        """
        Get patterns based on intake phase
        Maps intake phases to priorities and content types
        """
        phase_mapping = {
            'greeting': {'priority': ['NORMAL'], 'content_type': ['interview_question']},
            'symptom_discovery': {'priority': ['HIGH', 'NORMAL'], 'content_type': ['interview_question', 'clinical_clue']},
            'red_flags': {'priority': ['CRITICAL'], 'content_type': ['red_flag']},
            'differentials': {'priority': ['HIGH'], 'content_type': ['differential']},
            'context': {'priority': ['NORMAL', 'LOW'], 'content_type': ['interview_question']}
        }
        
        filters = phase_mapping.get(phase.lower(), {})
        priorities = filters.get('priority', [])
        content_types = filters.get('content_type', [])
        
        results = []
        for q in self.questions:
            if q.get('priority') in priorities and q.get('content_type') in content_types:
                results.append(q)
        
        return results
    
    def export_to_json(self, output_path: str = 'question_book_data.json'):
        """Export parsed patterns to JSON for easier access"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.questions, f, indent=2, ensure_ascii=False)
        logger.info(f"Exported {len(self.questions)} patterns to {output_path}")


# Test the RAG system
if __name__ == '__main__':
    rag = QuestionBookRAG()
    
    print("\n" + "="*80)
    print("RAG SYSTEM TEST")
    print("="*80)
    
    # Test searches with new schema
    print("\n1. Red flag patterns:")
    red_flags = rag.search_by_content_type('red_flag')
    print(f"   Found {len(red_flags)} red flags")
    if red_flags:
        sample = red_flags[0]
        print(f"   Sample: {sample.get('bot_question')}")
        print(f"   Domain: {sample.get('medical_domain')}")
        print(f"   Priority: {sample.get('priority')}")
    
    print("\n2. Clinical differentials:")
    diffs = rag.search_by_content_type('differential')
    print(f"   Found {len(diffs)} differential patterns")
    if diffs:
        sample = diffs[0]
        print(f"   Sample: {sample.get('bot_question')}")
        print(f"   Context: {sample.get('clinical_context')[:100]}...")
    
    print("\n3. Interview questions:")
    questions = rag.search_by_content_type('interview_question')
    print(f"   Found {len(questions)} interview questions")
    if questions:
        sample = questions[0]
        print(f"   Sample: {sample.get('bot_question')}")
        print(f"   Responses: {sample.get('expected_patient_responses', [])[:2]}")
    
    print("\n4. Priority-based retrieval test:")
    critical_patterns = rag.search_by_priority('CRITICAL')
    print(f"   Found {len(critical_patterns)} CRITICAL priority patterns")
    
    print("\n5. Domain-specific search (AbdominalPain):")
    abdominal = rag.search_by_system('AbdominalPain')
    print(f"   Found {len(abdominal)} patterns for AbdominalPain")
    
    # Export to JSON
    rag.export_to_json()
    
    print("\n" + "="*80)
    print("Test complete!")
    print("="*80)
