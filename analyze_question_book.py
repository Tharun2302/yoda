"""Analyze the Question Book structure"""
import docx
import json

doc = docx.Document('docx/Question BOOK.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}\n")
print("=" * 80)
print("DOCUMENT STRUCTURE ANALYSIS")
print("=" * 80)

# Find sections and question patterns
sections = []
questions = []
current_section = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Detect section headers (systems)
    if any(system in text for system in ['System', 'Cardiac', 'Respiratory', 'GI', 'Neurologic', 
                                         'Musculoskeletal', 'GU', 'Dermatologic', 'Endocrine', 'ENT']):
        if len(text) < 100 and text:  # Likely a header
            sections.append((i, text))
            current_section = text
    
    # Detect questions
    if 'Question:' in text or text.startswith('Q:') or text.startswith('Q.'):
        questions.append((i, current_section, text[:300]))
    
    # Sample first 100 non-empty paragraphs
    if i < 100 and text and len(text) > 10:
        print(f"Para {i}: {text[:150]}")

print("\n" + "=" * 80)
print(f"Found {len(sections)} potential sections")
print(f"Found {len(questions)} questions")
print("\nFirst 10 questions:")
for idx, (line, section, q) in enumerate(questions[:10]):
    print(f"\n{idx+1}. Line {line} [{section}]:")
    print(f"   {q}")

